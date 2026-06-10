from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document


def parse_docx(path: str | Path) -> tuple[str, dict]:
    """解析 DOCX 文件，返回 (提取文本, 元数据字典)"""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX 文件不存在: {path}")

    doc = Document(str(path))

    # 提取所有段落文本
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # 提取表格文本
    tables_text: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                tables_text.append(" | ".join(cells))

    full_text = "\n".join(paragraphs)
    if tables_text:
        full_text += "\n\n[表格内容]\n" + "\n".join(tables_text)

    # 元数据
    metadata = {
        "paragraphs": len(paragraphs),
        "tables": len(doc.tables),
    }

    return full_text, metadata


def parse_docx_simple(path: str | Path) -> str:
    """仅返回 DOCX 文本（忽略元数据）"""
    text, _ = parse_docx(path)
    return text
